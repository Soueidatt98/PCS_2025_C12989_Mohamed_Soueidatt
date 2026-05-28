# -*- coding: utf-8 -*-
"""
extract_judgments_to_json_v16.py

Extract Mauritanian Arabic criminal judgments from Word documents into:
  1) one canonical JSON per document
  2) a master JSONL file (one line per document)
  3) an extraction report CSV for quality control
  4) an optional Label Studio JSONL import file

Supported inputs:
  - .docx (native via python-docx)
  - .doc, .docm, .dotx, .dotm, .rtf, .odt (via LibreOffice headless conversion if available)

Usage examples:
  python extract_judgments_to_json_v16.py \
    --input-dir "C:\\judgments" \
    --output-dir "C:\\judgments_out"

  python extract_judgments_to_json_v16.py \
    --input-dir ./input \
    --output-dir ./out \
    --recursive \
    --export-labelstudio

Notes:
  - For legacy .doc files, LibreOffice/soffice must be installed and accessible.
  - The script preserves Arabic/RTL text as Unicode. JSON is written as UTF-8 with BOM
    to be friendlier with Windows editors.
"""

from __future__ import annotations

import argparse
import csv
import difflib
import io
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import uuid
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Iterator, List, Optional, Tuple

# ---------- Console encoding (Windows-friendly) ----------
try:
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

# ---------- 3rd-party dependency ----------
try:
    from docx import Document
    from docx.document import Document as _Document
    from docx.table import Table, _Cell
    from docx.text.paragraph import Paragraph
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
except Exception:
    print("ERROR: python-docx is required. Install with: pip install python-docx", file=sys.stderr)
    raise

# ---------- Constants ----------
NATIVE_DOCX_EXTS = {".docx"}
CONVERTIBLE_EXTS = {".doc", ".docm", ".dotx", ".dotm", ".rtf", ".odt"}
SUPPORTED_EXTS = NATIVE_DOCX_EXTS | CONVERTIBLE_EXTS

ARABIC_DIGIT_MAP = str.maketrans("٠١٢٣٤٥٦٧٨٩", "0123456789")
EASTERN_ARABIC_DIGIT_MAP = str.maketrans("۰۱۲۳۴۵۶۷۸۹", "0123456789")
TATWEEL = "ـ"
HARAKAT_RE = re.compile(r"[\u0617-\u061A\u064B-\u0652\u0670\u06D6-\u06ED]")
FORMAT_CHARS_RE = re.compile(r"[\u200e\u200f\u202a-\u202e\u2066-\u2069]")


ORDINAL_WORDS_RE = r"(?:أولا|اولا|أولاً|اولا|ثانيا|ثانياً|ثالثا|ثالثاً|رابعا|رابعاً|خامسا|خامساً|سادسا|سادساً|سابعا|سابعاً|ثامنا|ثامناً)"
SECTION_PREFIX_RE = r"(?ms)(^|[\n]|[ \t]{2,}|[\.؛،:]+\s*)(\s*(?:(?:" + ORDINAL_WORDS_RE + r")\s*[:：]?\s*)?)"
PROCEDURES_VARIANTS_RE = r"(?:الإجراءات|الاجراءات|الإجرءات|الاجرءات|الإجرائات|الاجرائات)"
FACTS_VARIANTS_RE = r"(?:الوقائع)"
REASONS_VARIANTS_RE = r"(?:الأسباب\s*والحيثيات|الاسباب\s*والحيثيات|أسباب\s*والحيثيات|اسباب\s*والحيثيات|الحيثيات\s*والأسباب|الحيثيات\s*والاسباب|حيثيات\s*وأسباب|حيثيات\s*واسباب|أسباب\s+الحكم|اسباب\s+الحكم|الأسباب|الاسباب|الحيثيات)"
LEGAL_BASIS_VARIANTS_RE = r"(?:لهذه\s+الأسباب|وعملا\s+وتأسيسا\s+على\s+ما\s*سبق(?:\s+وتطبيقا\s+للمواد)?|عملا\s+وتأسيسا\s+على\s+ما\s*سبق(?:\s+وتطبيقا\s+للمواد)?|تأسيسا\s+على\s+ما\s+سبق(?:\s+وتطبيقا\s+للمواد)?|وعملا\s+بالمواد|عملا\s+بالمواد|تطبيقا\s+للمواد|وبناء\s+على\s+المواد)"
OPERATIVE_VARIANTS_RE = r"(?:منطوق\s+الحكم|المنطوق)"

def make_section_title_pattern(title_regex: str) -> re.Pattern:
    # Require a strong title boundary after the heading word/phrase:
    # colon/punctuation, newline, or end of string.
    # Do NOT allow a generic space+body fallback here, otherwise ordinary body text like
    # "... ملخصا عن / الوقائع وطالب ..." can be misread as a new heading.
    return re.compile(SECTION_PREFIX_RE + r"(" + title_regex + r")\s*(?:[:：،,؛.]|(?=\n)|(?=$))")

SECTION_DEFINITIONS = [
    {
        "section_type": "facts_procedures",
        "canonical_title": "الوقائع والإجراءات",
        "required": False,
        "logical_types": ["facts", "procedures"],
        "patterns": [
            make_section_title_pattern(FACTS_VARIANTS_RE + r"\s*(?:و|/|\\)\s*" + PROCEDURES_VARIANTS_RE),
            make_section_title_pattern(PROCEDURES_VARIANTS_RE + r"\s*(?:و|/|\\)\s*" + FACTS_VARIANTS_RE),
        ],
    },
    {
        "section_type": "facts",
        "canonical_title": "الوقائع",
        "required": True,
        "logical_types": ["facts"],
        "patterns": [
            make_section_title_pattern(FACTS_VARIANTS_RE + r"(?!\s*(?:و|/|\\)\s*" + PROCEDURES_VARIANTS_RE + r")"),
        ],
    },
    {
        "section_type": "procedures",
        "canonical_title": "الإجراءات",
        "required": True,
        "logical_types": ["procedures"],
        "patterns": [
            make_section_title_pattern(PROCEDURES_VARIANTS_RE + r"(?!\s*(?:و|/|\\)\s*" + FACTS_VARIANTS_RE + r")"),
        ],
    },
    {
        "section_type": "court",
        "canonical_title": "المحكمة",
        "required": False,
        "logical_types": ["court"],
        "patterns": [
            make_section_title_pattern(r"المحكمة"),
        ],
    },
    {
        "section_type": "reasons",
        "canonical_title": "الأسباب",
        "required": True,
        "logical_types": ["reasons"],
        "patterns": [
            make_section_title_pattern(REASONS_VARIANTS_RE),
        ],
    },
    {
        "section_type": "legal_basis",
        "canonical_title": "لهذه الأسباب",
        "required": True,
        "logical_types": ["legal_basis"],
        "patterns": [
            make_section_title_pattern(LEGAL_BASIS_VARIANTS_RE),
        ],
    },
    {
        "section_type": "operative_part",
        "canonical_title": "منطوق الحكم",
        "required": True,
        "logical_types": ["operative_part"],
        "patterns": [
            make_section_title_pattern(OPERATIVE_VARIANTS_RE),
        ],
    },
]

SECTION_TITLES = {item["section_type"]: item["canonical_title"] for item in SECTION_DEFINITIONS}

FUZZY_HEADING_SYNONYMS = {
    "facts_procedures": [
        "الوقائع والإجراءات", "الوقائع والاجراءات", "الوقائع والإجرءات", "الوقائع والاجرءات",
        "الإجراءات والوقائع", "الاجراءات والوقائع", "الإجرءات والوقائع", "الاجرءات والوقائع",
    ],
    "facts": ["الوقائع"],
    "procedures": ["الإجراءات", "الاجراءات", "الإجرءات", "الاجرءات", "الإجرائات", "الاجرائات"],
    "court": ["المحكمة"],
    "reasons": [
        "الأسباب", "الاسباب", "أسباب الحكم", "اسباب الحكم",
        "الأسباب والحيثيات", "الاسباب والحيثيات", "أسباب والحيثيات", "اسباب والحيثيات",
        "الحيثيات والأسباب", "الحيثيات والاسباب", "حيثيات وأسباب", "حيثيات واسباب",
        "الحيثيات"
    ],
    "legal_basis": [
        "لهذه الأسباب", "لهذه الاسباب", "وعملا بالمواد", "عملا بالمواد",
        "وبناء على المواد", "بناء على المواد", "وتطبيقا للمواد", "تطبيقا للمواد",
        "تأسيسا على ما سبق", "تاسيسا على ما سبق", "وعملا وتأسيسا على ما سبق", "عملا وتأسيسا على ما سبق"
    ],
    "operative_part": ["منطوق الحكم", "المنطوق"],
}

LABEL_VARIANTS = {
    "case_number": ["رقم القضية", "رقم النيابة", "الملف رقم", "رقم الملف", "القضية رقم"],
    "judgment_number": ["الحكم رقم", "رقم الحكم", "منطوق الحكم رقم", "الحكمين رقمي", "الحكمان رقما"],
    "deposit_date": ["تاريخ الإيداع"],
    "judgment_date": ["تاريخه", "تاريخ الحكم", "بتاريخ", "تاريخي"],
    "charge": ["التهمة", "التهمه"],
    "appearance_type": ["وصفه"],
    "degree": ["درجته"],
    "summary": ["ملخصه"],
    "defendants_header": ["المتهمين", "المتهم", "المتهمان", "المتهمون"],
    "civil_parties_header": ["الاطراف المدنية", "الأطراف المدنية", "الأطراف المدني", "الطرف المدني", "الاطراف"],
}

ROLE_MAP = {
    "defendants_header": "defendant",
    "civil_parties_header": "civil_party",
}

# ---------- Data structures ----------
@dataclass
class ExtractionResult:
    data: dict
    report_row: dict


# ---------- Utility functions ----------
def normalize_digits(s: str) -> str:
    return s.translate(ARABIC_DIGIT_MAP).translate(EASTERN_ARABIC_DIGIT_MAP)


def basic_normalize(s: str, remove_diacritics: bool = False) -> str:
    if s is None:
        return ""
    s = s.replace("\ufeff", " ")
    s = FORMAT_CHARS_RE.sub("", s)
    s = s.replace(TATWEEL, "")
    s = s.replace("\xa0", " ")
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = s.replace("–", "-").replace("—", "-").replace("：", ":")
    s = normalize_digits(s)
    if remove_diacritics:
        s = HARAKAT_RE.sub("", s)
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r" ?\n ?", "\n", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()


def slugify_filename(s: str, max_len: int = 120) -> str:
    s = basic_normalize(s)
    s = re.sub(r"[\\/:*?\"<>|]+", "_", s)
    s = re.sub(r"\s+", "_", s)
    s = s.strip("._")
    if not s:
        s = str(uuid.uuid4())
    return s[:max_len]


def iso_date_or_none(s: Optional[str]) -> Optional[str]:
    if not s:
        return None
    s = normalize_digits(s)
    m = re.search(r"\b([0-3]?\d)\s*[/\\-]\s*([01]?\d)\s*[/\\-]\s*((?:19|20)?\d{2})\b", s)
    if not m:
        return None
    day, month, year = m.groups()
    day = int(day)
    month = int(month)
    year = int(year)
    if year < 100:
        year += 2000
    if 1 <= month <= 12 and 1 <= day <= 31:
        return f"{year:04d}-{month:02d}-{day:02d}"
    return None


def extract_all_dates(text: str) -> List[str]:
    vals: List[str] = []
    for m in re.finditer(r"\b([0-3]?\d)\s*[/\\-]\s*([01]?\d)\s*[/\\-]\s*((?:19|20)?\d{2})\b", normalize_digits(text)):
        iso = iso_date_or_none(m.group(0))
        if iso and iso not in vals:
            vals.append(iso)
    return vals


def is_valid_docx(path: Path) -> bool:
    try:
        return zipfile.is_zipfile(str(path))
    except Exception:
        return False


def is_temporary_office_file(path: Path) -> bool:
    """Ignore Word/LibreOffice lock and temp files.

    Examples:
      - ~$document.docx  (Word owner/lock file)
      - .~lock.document#  (LibreOffice lock file)
      - temporary files ending in .tmp
    """
    name = path.name
    lname = name.lower()
    return (
        name.startswith('~$')
        or lname.startswith('.~lock')
        or lname.endswith('.tmp')
        or lname.startswith('~wr')
    )


def can_open_with_python_docx(path: Path) -> Tuple[bool, Optional[str]]:
    try:
        Document(str(path))
        return True, None
    except Exception as e:
        return False, str(e)


def find_soffice() -> Optional[str]:
    candidates = [
        os.environ.get("SOFFICE_PATH"),
        shutil.which("soffice"),
        shutil.which("libreoffice"),
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    for c in candidates:
        if c and Path(c).exists():
            return c
    return None


def convert_to_docx(input_path: Path, temp_dir: Path) -> Tuple[Optional[Path], Optional[str]]:
    soffice = find_soffice()
    if not soffice:
        return None, "LibreOffice/soffice not found. Needed to convert non-.docx files."

    cmd = [soffice, "--headless", "--convert-to", "docx", "--outdir", str(temp_dir), str(input_path)]
    try:
        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
    except Exception as e:
        return None, f"LibreOffice conversion failed to start: {e}"

    if proc.returncode != 0:
        stderr = (proc.stderr or "").strip()
        stdout = (proc.stdout or "").strip()
        return None, f"LibreOffice conversion failed (code {proc.returncode}). stdout={stdout} stderr={stderr}"

    out_path = temp_dir / f"{input_path.stem}.docx"
    if out_path.exists() and is_valid_docx(out_path):
        return out_path, None

    generated = list(temp_dir.glob("*.docx"))
    for cand in generated:
        if is_valid_docx(cand):
            return cand, None
    return None, "LibreOffice reported success but no valid .docx output was found."


# ---------- DOCX reading preserving body order ----------
def iter_block_items(parent) -> Iterator[Paragraph | Table]:
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Unsupported parent for iter_block_items")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def cell_text(cell: _Cell) -> str:
    parts: List[str] = []
    for block in iter_block_items(cell):
        if isinstance(block, Paragraph):
            t = block.text.strip()
            if t:
                parts.append(t)
        elif isinstance(block, Table):
            for row in block.rows:
                row_text = " | ".join(cell_text(c).strip() for c in row.cells if cell_text(c).strip())
                if row_text:
                    parts.append(row_text)
    return "\n".join(parts)


def read_docx_parts(path: Path) -> dict:
    doc = Document(str(path))

    body_parts: List[str] = []
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            t = block.text.strip()
            if t:
                body_parts.append(t)
        elif isinstance(block, Table):
            for row in block.rows:
                row_parts = []
                for c in row.cells:
                    txt = cell_text(c).strip()
                    if txt:
                        row_parts.append(txt)
                if row_parts:
                    body_parts.append(" | ".join(row_parts))

    header_parts: List[str] = []
    footer_parts: List[str] = []
    for sect in doc.sections:
        if sect.header:
            for p in sect.header.paragraphs:
                t = p.text.strip()
                if t:
                    header_parts.append(t)
        if sect.footer:
            for p in sect.footer.paragraphs:
                t = p.text.strip()
                if t:
                    footer_parts.append(t)

    return {
        "body_text_raw": "\n".join(body_parts),
        "header_text_aux": "\n".join(header_parts),
        "footer_text_aux": "\n".join(footer_parts),
    }


# ---------- Section extraction ----------
def _normalize_heading_candidate(s: str) -> str:
    s = basic_normalize(s, remove_diacritics=True)
    s = s.replace("أ", "ا").replace("إ", "ا").replace("آ", "ا")
    s = s.replace("ؤ", "و").replace("ئ", "ي")
    s = re.sub(r"\b(?:" + ORDINAL_WORDS_RE + r")\b", " ", s)
    s = re.sub(r"^[\s\-–—ـ:：؛،,.()\[\]/\\]+|[\s\-–—ـ:：؛،,.()\[\]/\\]+$", "", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def _section_compare_key(s: str) -> str:
    s = _normalize_heading_candidate(s)
    toks = []
    for tok in s.split():
        tok = re.sub(r"^ال", "", tok)
        toks.append(tok)
    return " ".join(toks).strip()


def _classify_heading_fuzzy(raw_heading: str, allowed_types: Optional[set] = None) -> Optional[Tuple[str, str, List[str], float]]:
    cand = _normalize_heading_candidate(raw_heading)
    if not cand:
        return None

    # Hard gate: only try fuzzy classification on short heading-like chunks.
    if len(cand) > 28 or len(cand.split()) > 4:
        return None

    cand_key = _section_compare_key(cand)
    best = None
    best_ratio = 0.0
    for stype, syns in FUZZY_HEADING_SYNONYMS.items():
        if allowed_types and stype not in allowed_types:
            continue
        for syn in syns:
            syn_n = _normalize_heading_candidate(syn)
            syn_key = _section_compare_key(syn_n)
            ratio = difflib.SequenceMatcher(None, cand_key, syn_key).ratio()
            if cand_key == syn_key:
                ratio = 1.0
            if ratio > best_ratio:
                best_ratio = ratio
                best = (stype, syn)

    if not best:
        return None

    # Controlled thresholds: allow moderate fuzziness only for known tiny title vocabulary.
    threshold = 0.90
    if best[0] in {"facts_procedures", "reasons", "operative_part"}:
        threshold = 0.86
    if best_ratio < threshold:
        return None

    stype, syn = best
    logical = next((d.get("logical_types", [d["section_type"]]) for d in SECTION_DEFINITIONS if d["section_type"] == stype), [stype])
    canon = SECTION_TITLES.get(stype, syn)
    return stype, canon, logical, best_ratio


def _find_fuzzy_section_candidates(clean_text: str, existing_matches: List[dict], allowed_types: Optional[set] = None) -> List[dict]:
    # Only search around real heading boundaries; avoid scanning arbitrary body text.
    candidates = []
    occupied = [(m["start"], m["end_title"]) for m in existing_matches]
    boundary_re = re.compile(r"(?ms)(^|\n|[ \t]{2,}|[\.؛،:]+\s*)(.{1,60}?)(?=(?:[:：،,؛.]|\n|$))")
    keyword_hint = re.compile(r"(وقع|اجر|جرء|سبب|حيث|منطوق|محكم|مواد|تطبيقا|تاسيسا|لهذه)")

    for m in boundary_re.finditer(clean_text):
        frag = (m.group(2) or "").strip()
        if not frag:
            continue
        start = m.start(2)
        end = m.end(2)
        # avoid overlaps with already found explicit headings
        if any(not (end <= a or start >= b) for a, b in occupied):
            continue
        frag_n = _normalize_heading_candidate(frag)
        if not frag_n or not keyword_hint.search(frag_n):
            continue
        classified = _classify_heading_fuzzy(frag, allowed_types=allowed_types)
        if not classified:
            continue
        stype, canon, logical, score = classified
        candidates.append({
            "section_type": stype,
            "title_original": frag,
            "title_canonical": canon,
            "required": next((d["required"] for d in SECTION_DEFINITIONS if d["section_type"] == stype), False),
            "logical_types": logical,
            "start": start,
            "end_title": end,
            "match_len": end - start,
            "fuzzy_score": score,
        })
    return candidates


def find_sections(clean_text: str) -> Tuple[List[dict], List[str], str]:
    warnings: List[str] = []
    candidates = []
    seen = set()

    for definition in SECTION_DEFINITIONS:
        for pattern in definition["patterns"]:
            for m in pattern.finditer(clean_text):
                start, end = m.start(), m.end()
                key = (definition["section_type"], start, end)
                if key in seen:
                    continue
                seen.add(key)
                title_start = m.start(3) if m.lastindex and m.lastindex >= 3 else m.start()
                title_end = m.end()
                title_original = clean_text[title_start:title_end].strip() if title_start < title_end else m.group(0).strip()
                candidates.append({
                    "section_type": definition["section_type"],
                    "title_original": title_original,
                    "title_canonical": definition["canonical_title"],
                    "required": definition["required"],
                    "logical_types": definition.get("logical_types", [definition["section_type"]]),
                    "start": title_start,
                    "end_title": title_end,
                    "match_len": title_end - title_start,
                })

    candidates.sort(key=lambda x: (x["start"], -x["match_len"]))

    matches: List[dict] = []
    for cand in candidates:
        if not matches:
            matches.append(cand)
            continue
        prev = matches[-1]
        if cand["start"] >= prev["end_title"]:
            matches.append(cand)
        else:
            if cand["match_len"] > prev["match_len"]:
                matches[-1] = cand

    # Drop an initial false-positive "court" match from the document header
    if matches and matches[0]["section_type"] == "court" and matches[0]["start"] < max(200, len(clean_text) // 20):
        if not re.search(r"[:：]\s*$", matches[0]["title_original"]):
            matches = matches[1:]

    found_types = {m["section_type"] for m in matches}
    found_logical_types = set()
    for m in matches:
        found_logical_types.update(m.get("logical_types", [m["section_type"]]))

    # Controlled fuzzy rescue for missing section headings (typos/orthographic variants).
    missing_required = []
    for definition in SECTION_DEFINITIONS:
        logical_types = definition.get("logical_types", [definition["section_type"]])
        if definition["required"] and not all(t in found_logical_types for t in logical_types):
            missing_required.append(definition["section_type"])

    if missing_required:
        allowed_types = {t for t in missing_required if t != "court"}
        fuzzy_candidates = _find_fuzzy_section_candidates(clean_text, matches, allowed_types=allowed_types)
        if fuzzy_candidates:
            matches.extend(fuzzy_candidates)
            matches.sort(key=lambda x: (x["start"], -x["match_len"]))
            collapsed = []
            for cand in matches:
                if not collapsed:
                    collapsed.append(cand)
                    continue
                prev = collapsed[-1]
                if cand["start"] >= prev["end_title"]:
                    collapsed.append(cand)
                else:
                    prev_score = prev.get("fuzzy_score", 1.0)
                    cand_score = cand.get("fuzzy_score", 1.0)
                    prev_explicit = "fuzzy_score" not in prev
                    cand_explicit = "fuzzy_score" not in cand
                    if cand_explicit and not prev_explicit:
                        collapsed[-1] = cand
                    elif cand_explicit == prev_explicit and (cand_score > prev_score or cand["match_len"] > prev["match_len"]):
                        collapsed[-1] = cand
            matches = collapsed
            found_logical_types = set()
            for m in matches:
                found_logical_types.update(m.get("logical_types", [m["section_type"]]))

    # If no explicit reasons heading exists but there is a "court" heading between facts/procedures and legal basis,
    # treat the last such court block as reasons. This matches files where "المحكمة" contains the reasoning.
    if "reasons" not in found_logical_types:
        legal_start = next((m["start"] for m in matches if m["section_type"] == "legal_basis"), None)
        candidate_idx = None
        for idx, m in enumerate(matches):
            if m["section_type"] == "court" and (legal_start is None or m["start"] < legal_start):
                candidate_idx = idx
        if candidate_idx is not None:
            matches[candidate_idx]["section_type"] = "reasons"
            matches[candidate_idx]["title_canonical"] = "الأسباب"
            matches[candidate_idx]["logical_types"] = ["reasons"]
            found_logical_types.add("reasons")

    # Merge consecutive duplicate headings of the same logical section when the later one is just
    # a sub-heading marker (common for legal basis: "لهذه الأسباب" then "وبناء على المواد...").
    merged_matches: List[dict] = []
    for m in matches:
        if not merged_matches:
            merged_matches.append(m)
            continue
        prev = merged_matches[-1]
        gap = clean_text[prev["end_title"]:m["start"]]
        same_logical = prev["section_type"] == m["section_type"]
        trivial_gap = not gap.strip() or re.fullmatch(r"[\s\r\n\t\-–—ـ:：؛،,.()\[\]/\\]*", gap or "") is not None
        if same_logical and trivial_gap:
            prev_title = prev.get("title_original", "").strip()
            cur_title = m.get("title_original", "").strip()
            if cur_title and cur_title not in prev_title:
                prev["title_original"] = (prev_title + " / " + cur_title).strip(" /")
            prev["end_title"] = m["end_title"]
            prev["match_len"] = prev["end_title"] - prev["start"]
        else:
            merged_matches.append(m)
    matches = merged_matches

    # Recompute found logical types after merge.
    found_logical_types = set()
    for m in matches:
        found_logical_types.update(m.get("logical_types", [m["section_type"]]))

    fuzzy_used = [m for m in matches if "fuzzy_score" in m]
    for m in fuzzy_used:
        warnings.append(f"fuzzy_heading_used:{m['section_type']}:{m.get('title_original','')}:{m.get('fuzzy_score',0):.2f}")

    for definition in SECTION_DEFINITIONS:
        logical_types = definition.get("logical_types", [definition["section_type"]])
        if definition["required"] and not all(t in found_logical_types for t in logical_types):
            warnings.append(f"Section heading not found: {definition['section_type']}")

    sections: List[dict] = []
    header_text = clean_text[: matches[0]["start"]].strip() if matches else clean_text.strip()

    for i, item in enumerate(matches):
        next_start = matches[i + 1]["start"] if i + 1 < len(matches) else len(clean_text)
        section_text = clean_text[item["end_title"]: next_start].strip()
        sections.append({
            "section_id": f"S{i+1}",
            "section_type": item["section_type"],
            "logical_section_types": item.get("logical_types", [item["section_type"]]),
            "title_original": item["title_original"],
            "title_canonical": item["title_canonical"],
            "order": i + 1,
            "text": section_text,
            "char_start": item["start"],
            "char_end": next_start,
        })

    return sections, warnings, header_text


# ---------- Metadata extraction ----------
def extract_court(text: str) -> Optional[str]:
    for line in text.splitlines():
        line = line.strip()
        if line.startswith("محكمة"):
            return line.split("|")[0].strip()
    m = re.search(r"(محكمة[^\n]+)", text)
    return m.group(1).split("|")[0].strip() if m else None


def extract_chamber(text: str) -> Optional[str]:
    for line in text.splitlines():
        line = line.strip()
        if line.startswith("الغرفة") or line.startswith("غرفة"):
            return line.split("|")[0].strip()
    m = re.search(r"((?:الغرفة|غرفة)[^\n]+)", text)
    return m.group(1).split("|")[0].strip() if m else None


def extract_label_value_multiline(text: str, labels: List[str], stop_labels: Optional[List[str]] = None) -> Optional[str]:
    stop_labels = stop_labels or []
    all_stops = sorted(
        set(stop_labels + [v for vals in LABEL_VARIANTS.values() for v in vals] + [v for v in SECTION_TITLES.values()]),
        key=len,
        reverse=True,
    )
    label_pattern = "|".join(re.escape(lbl) for lbl in sorted(labels, key=len, reverse=True))
    stop_pattern = "|".join(re.escape(lbl) for lbl in all_stops)

    pat = re.compile(rf"(?ms)^\s*(?:{label_pattern})\s*[:：]?\s*(.*?)\s*(?=^\s*(?:{stop_pattern})\s*[:：]?|\Z)")
    m = pat.search(text)
    if not m:
        return None
    val = m.group(1).strip()
    val = re.sub(r"\n{2,}", "\n", val)
    val = re.sub(r"[ \t]{2,}", " ", val)
    return val or None


def extract_case_number(text: str) -> Optional[str]:
    patterns = [
        r"(?:رقم\s*القضية|القضية\s*رقم)\s*[:：]?\s*([0-9]+\s*[/\\]\s*[0-9]{2,4})",
        r"(?:رقم\s*النيابة)\s*[:：]?\s*([0-9]+\s*[/\\]\s*[0-9]{2,4})",
        r"(?:الملف\s*رقم|رقم\s*الملف)\s*[:：]?\s*([0-9]+\s*[/\\]\s*[0-9]{2,4})",
    ]
    txt = normalize_digits(text)
    for pat in patterns:
        m = re.search(pat, txt)
        if m:
            return re.sub(r"\s*[/\\]\s*", "/", m.group(1)).strip()
    return None


def extract_judgment_number(text: str) -> Optional[str]:
    txt = normalize_digits(text)
    m = re.search(r"(?:الحكم\s*رقم|رقم\s*الحكم|منطوق\s*الحكم\s*رقم)\s*[:：]?\s*([0-9]+\s*[/\\]\s*[0-9]{2,4})", txt)
    if m:
        return re.sub(r"\s*[/\\]\s*", "/", m.group(1)).strip()
    return None


def extract_hearing_dates(text: str) -> List[str]:
    txt = normalize_digits(text)
    candidates: List[str] = []

    for m in re.finditer(r"([0-3]?\d)\s*و\s*([0-3]?\d)\s*[/\\-]\s*([01]?\d)\s*[/\\-]\s*((?:19|20)?\d{2})", txt):
        d1, d2, month, year = m.groups()
        for day in (d1, d2):
            iso = iso_date_or_none(f"{day}/{month}/{year}")
            if iso and iso not in candidates:
                candidates.append(iso)

    patterns = [
        r"الجلسة\s+المنعقدة\s+بتاريخ\s*[:：]?\s*([0-3]?\d\s*[/\\-]\s*[01]?\d\s*[/\\-]\s*(?:19|20)?\d{2})",
        r"يوم[^\n]{0,80}?الموافق(?:ين)?\s*[:：]?\s*([0-3]?\d\s*[/\\-]\s*[01]?\d\s*[/\\-]\s*(?:19|20)?\d{2})",
        r"بتاريخ(?:ي)?\s*[:：]?\s*([0-3]?\d\s*[/\\-]\s*[01]?\d\s*[/\\-]\s*(?:19|20)?\d{2})",
    ]
    for pat in patterns:
        for m in re.finditer(pat, txt):
            iso = iso_date_or_none(m.group(1))
            if iso and iso not in candidates:
                candidates.append(iso)
    return candidates


def split_header_party_value(raw_value: Optional[str]) -> List[dict]:
    """
    Conservative but practical splitting:
      - inline numbered items such as:
          "1- ابراهيم ... 2 سيداتي ... 3 الداه ..."
        are split into separate resolved entries
      - bullet/numbered lines are split into separate resolved entries
      - group markers like "وأخريات" remain unresolved
      - otherwise keep as one resolved entry
    """
    if not raw_value:
        return []
    v = raw_value.strip()
    if not v:
        return []

    # Remove accidental leading punctuation or separators left after label extraction.
    v = re.sub(r"^[\s:：\-ـ]+", "", v)

    if re.search(r"\b(?:وأخريات|وآخرون|وآخرين|وغيرهم|وأخرى|وأخريان)\b", v):
        return [{"raw_text": v, "resolved": False}]

    inline = re.sub(r"\s+", " ", v).strip()
    if re.search(r"^(?:\d{1,2}\s*(?:[\-ـ\.)])?\s+)", inline) or re.search(r"\s\d{1,2}\s*(?:[\-ـ\.)])?\s+", inline):
        parts = re.split(r"\s*(?:(?<=^)|(?<=\s))\d{1,2}\s*(?:[\-ـ\.)])?\s+", inline)
        parts = [p.strip(" ،,;؛.-") for p in parts if p.strip(" ،,;؛.-")]
        # Only trust this split if it actually yields multiple substantial items.
        if len(parts) >= 2:
            return [{"raw_text": item, "resolved": True} for item in parts]

    split_lines = []
    for ln in v.splitlines():
        ln2 = ln.strip()
        m = re.match(r"^(?:[\-/*•]|\d+\s*[\.)\-ـ])\s*(.+)$", ln2)
        if m:
            split_lines.append(m.group(1).strip(" ،,;؛."))
    if split_lines:
        return [{"raw_text": item, "resolved": True} for item in split_lines if item]

    nonempty_lines = [ln.strip() for ln in v.splitlines() if ln.strip()]
    if len(nonempty_lines) > 1:
        return [{"raw_text": item, "resolved": True} for item in nonempty_lines]

    if "،" in v:
        chunks = [c.strip() for c in v.split("،") if c.strip()]
        if len(chunks) > 1:
            return [{"raw_text": c, "resolved": True} for c in chunks]

    return [{"raw_text": v, "resolved": True}]


def build_participants_index(header_entities: dict) -> List[dict]:
    participants = []
    n = 1
    for field_name, entries in header_entities.items():
        role = ROLE_MAP.get(field_name)
        if not role:
            continue
        for entry in entries:
            pid = f"P{n}"
            participants.append({
                "participant_id": pid,
                "role": role,
                "is_group": not entry.get("resolved", False),
                "display_name": entry.get("raw_text"),
                "raw_mentions": [entry.get("raw_text")],
                "from_header": True,
                "attributes": {},
            })
            n += 1
    return participants


# ---------- Core extraction ----------
def extract_document(input_path: Path) -> ExtractionResult:
    warnings: List[str] = []
    errors: List[str] = []
    conversion_used = False
    missing_fields: List[str] = []

    work_path = input_path
    temp_dir_obj = None

    try:
        ext = input_path.suffix.lower()
        if ext not in SUPPORTED_EXTS:
            raise ValueError(f"Unsupported extension: {ext}")

        if ext in CONVERTIBLE_EXTS:
            temp_dir_obj = tempfile.TemporaryDirectory(prefix="judgment_convert_")
            converted, err = convert_to_docx(input_path, Path(temp_dir_obj.name))
            if err:
                raise RuntimeError(err)
            if not converted:
                raise RuntimeError("Conversion failed without output path.")
            work_path = converted
            conversion_used = True

        # Prefer actual openability over a zip-only check. Some files can fail a
        # superficial validation step or need a fallback conversion/re-save.
        can_open, open_err = can_open_with_python_docx(work_path)
        if not can_open:
            # If the original extension looked native but python-docx cannot open it,
            # try a LibreOffice re-save to a fresh .docx before failing.
            if ext in NATIVE_DOCX_EXTS and not conversion_used:
                temp_dir_obj = tempfile.TemporaryDirectory(prefix="judgment_resave_")
                converted, err = convert_to_docx(input_path, Path(temp_dir_obj.name))
                if converted:
                    work_path = converted
                    conversion_used = True
                    can_open, open_err = can_open_with_python_docx(work_path)
                elif err:
                    warnings.append(f"native_docx_resave_failed: {err}")
            if not can_open:
                zip_ok = is_valid_docx(work_path)
                raise RuntimeError(f"Could not open as .docx after conversion/open step. zip_ok={zip_ok}; open_error={open_err}")

        parts = read_docx_parts(work_path)
        body_text_raw = parts["body_text_raw"]
        header_aux_raw = parts["header_text_aux"]
        footer_aux_raw = parts["footer_text_aux"]

        full_text_raw = body_text_raw.strip()
        full_text_clean = basic_normalize(full_text_raw, remove_diacritics=False)

        aux_text_combined = "\n".join(x for x in [header_aux_raw, full_text_raw, footer_aux_raw] if x and x.strip())
        aux_text_clean = basic_normalize(aux_text_combined, remove_diacritics=False)

        sections, section_warnings, header_text = find_sections(full_text_clean)
        warnings.extend(section_warnings)

        extraction_source = header_text if header_text else aux_text_clean
        fallback_source = aux_text_clean

        court = extract_court(extraction_source) or extract_court(fallback_source)
        chamber = extract_chamber(extraction_source) or extract_chamber(fallback_source)
        case_number = extract_case_number(extraction_source) or extract_case_number(fallback_source)
        judgment_number = extract_judgment_number(extraction_source) or extract_judgment_number(fallback_source)
        charge_raw = extract_label_value_multiline(extraction_source, LABEL_VARIANTS["charge"]) or extract_label_value_multiline(fallback_source, LABEL_VARIANTS["charge"])
        deposit_date_raw = extract_label_value_multiline(extraction_source, LABEL_VARIANTS["deposit_date"]) or extract_label_value_multiline(fallback_source, LABEL_VARIANTS["deposit_date"])
        judgment_date_raw = extract_label_value_multiline(extraction_source, LABEL_VARIANTS["judgment_date"]) or extract_label_value_multiline(fallback_source, LABEL_VARIANTS["judgment_date"])
        appearance_type = extract_label_value_multiline(extraction_source, LABEL_VARIANTS["appearance_type"]) or extract_label_value_multiline(fallback_source, LABEL_VARIANTS["appearance_type"])
        degree = extract_label_value_multiline(extraction_source, LABEL_VARIANTS["degree"]) or extract_label_value_multiline(fallback_source, LABEL_VARIANTS["degree"])
        summary = extract_label_value_multiline(extraction_source, LABEL_VARIANTS["summary"]) or extract_label_value_multiline(fallback_source, LABEL_VARIANTS["summary"])
        defendants_raw = extract_label_value_multiline(extraction_source, LABEL_VARIANTS["defendants_header"]) or extract_label_value_multiline(fallback_source, LABEL_VARIANTS["defendants_header"])
        civil_raw = extract_label_value_multiline(extraction_source, LABEL_VARIANTS["civil_parties_header"]) or extract_label_value_multiline(fallback_source, LABEL_VARIANTS["civil_parties_header"])

        charge_list = [charge_raw.strip(" .،؛;")] if charge_raw else []
        hearing_dates = extract_hearing_dates(full_text_clean)

        header_entities = {
            "defendants_header": split_header_party_value(defendants_raw),
            "civil_parties_header": split_header_party_value(civil_raw),
        }
        participants_index = build_participants_index(header_entities)

        if not court:
            missing_fields.append("court")
        if not chamber:
            missing_fields.append("chamber")
        if not case_number:
            missing_fields.append("case_number")
        if not judgment_number:
            missing_fields.append("judgment_number")
        if not charge_list:
            missing_fields.append("charge_list")
        if not sections:
            missing_fields.append("sections")

        stem = slugify_filename(input_path.stem)
        doc_id = stem

        data = {
            "schema_version": "1.6",
            "doc_id": doc_id,
            "source": {
                "file_name": input_path.name,
                "file_path": str(input_path),
                "file_format": input_path.suffix.lower().lstrip("."),
                "conversion_used": conversion_used,
                "converted_from": input_path.suffix.lower() if conversion_used else None,
            },
            "document_info": {
                "country": "MR",
                "language": "ar",
                "document_type": "judgment",
                "court": court,
                "chamber": chamber,
            },
            "metadata": {
                "case_number": case_number,
                "judgment_number": judgment_number,
                "deposit_date": iso_date_or_none(deposit_date_raw) or iso_date_or_none(full_text_clean),
                "judgment_date": iso_date_or_none(judgment_date_raw) or (extract_all_dates(header_text)[0] if header_text and extract_all_dates(header_text) else None),
                "hearing_dates": hearing_dates,
                "charge_list": charge_list,
                "appearance_type": appearance_type,
                "degree": degree,
                "summary": summary,
            },
            "header_entities": header_entities,
            "participants_index": participants_index,
            "header_text": header_text,
            "sections": sections,
            "text": {
                "full_text_raw": full_text_raw,
                "full_text_clean": full_text_clean,
                "header_text_aux": basic_normalize(header_aux_raw) if header_aux_raw else "",
                "footer_text_aux": basic_normalize(footer_aux_raw) if footer_aux_raw else "",
            },
            "annotations": {
                "ner_spans": [],
                "role_mentions": [],
                "coreference_links": [],
            },
            "processing": {
                "extraction_status": "ok" if not errors else "error",
                "normalization_status": "ok",
                "section_parser_version": "v16",
                "warnings": warnings,
                "errors": errors,
                "missing_fields": missing_fields,
            },
        }

        report_row = {
            "doc_id": doc_id,
            "source_file": input_path.name,
            "source_ext": input_path.suffix.lower(),
            "conversion_used": conversion_used,
            "status": "ok",
            "text_len": len(full_text_clean),
            "sections_found": len(sections),
            "missing_fields": "|".join(missing_fields),
            "warnings_count": len(warnings),
            "warnings": " | ".join(warnings),
            "errors": "",
        }

        return ExtractionResult(data=data, report_row=report_row)

    except Exception as e:
        errors.append(str(e))
        doc_id = slugify_filename(input_path.stem)
        data = {
            "schema_version": "1.6",
            "doc_id": doc_id,
            "source": {
                "file_name": input_path.name,
                "file_path": str(input_path),
                "file_format": input_path.suffix.lower().lstrip("."),
                "conversion_used": conversion_used,
                "converted_from": input_path.suffix.lower() if conversion_used else None,
            },
            "document_info": {
                "country": "MR",
                "language": "ar",
                "document_type": "judgment",
                "court": None,
                "chamber": None,
            },
            "metadata": {
                "case_number": None,
                "judgment_number": None,
                "deposit_date": None,
                "judgment_date": None,
                "hearing_dates": [],
                "charge_list": [],
                "appearance_type": None,
                "degree": None,
                "summary": None,
            },
            "header_entities": {
                "defendants_header": [],
                "civil_parties_header": [],
            },
            "participants_index": [],
            "header_text": "",
            "sections": [],
            "text": {
                "full_text_raw": "",
                "full_text_clean": "",
                "header_text_aux": "",
                "footer_text_aux": "",
            },
            "annotations": {
                "ner_spans": [],
                "role_mentions": [],
                "coreference_links": [],
            },
            "processing": {
                "extraction_status": "error",
                "normalization_status": "failed",
                "section_parser_version": "v16",
                "warnings": warnings,
                "errors": errors,
                "missing_fields": ["all"],
            },
        }
        report_row = {
            "doc_id": doc_id,
            "source_file": input_path.name,
            "source_ext": input_path.suffix.lower(),
            "conversion_used": conversion_used,
            "status": "error",
            "text_len": 0,
            "sections_found": 0,
            "missing_fields": "all",
            "warnings_count": len(warnings),
            "warnings": " | ".join(warnings),
            "errors": " | ".join(errors),
        }
        return ExtractionResult(data=data, report_row=report_row)
    finally:
        if temp_dir_obj is not None:
            temp_dir_obj.cleanup()


# ---------- File walking and writing ----------
def iter_input_files(input_dir: Path, recursive: bool) -> Iterator[Path]:
    iterator = input_dir.rglob("*") if recursive else input_dir.glob("*")
    for p in iterator:
        if not p.is_file():
            continue
        if is_temporary_office_file(p):
            continue
        if p.suffix.lower() in SUPPORTED_EXTS:
            yield p


def write_json(path: Path, obj: dict) -> None:
    path.write_text(json.dumps(obj, ensure_ascii=False, indent=2), encoding="utf-8-sig")


def write_jsonl(path: Path, items: Iterable[dict]) -> None:
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        for item in items:
            f.write(json.dumps(item, ensure_ascii=False) + "\n")


def write_csv(path: Path, rows: List[dict]) -> None:
    fieldnames = [
        "doc_id", "source_file", "source_ext", "conversion_used", "status",
        "text_len", "sections_found", "missing_fields", "warnings_count",
        "warnings", "errors"
    ]
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        for row in rows:
            writer.writerow(row)


def build_labelstudio_task(doc: dict) -> dict:
    return {
        "id": doc["doc_id"],
        "data": {
            "text": doc["text"]["full_text_clean"],
            "doc_id": doc["doc_id"],
            "case_number": doc["metadata"].get("case_number"),
            "judgment_number": doc["metadata"].get("judgment_number"),
            "court": doc["document_info"].get("court"),
            "chamber": doc["document_info"].get("chamber"),
            "charge_list": doc["metadata"].get("charge_list", []),
            "sections": [
                {
                    "section_id": s["section_id"],
                    "section_type": s["section_type"],
                    "title_original": s["title_original"],
                    "title_canonical": s.get("title_canonical"),
                    "char_start": s["char_start"],
                    "char_end": s["char_end"],
                }
                for s in doc.get("sections", [])
            ],
        },
        "meta": {
            "source_file": doc["source"]["file_name"],
            "missing_fields": doc["processing"].get("missing_fields", []),
            "warnings": doc["processing"].get("warnings", []),
        }
    }


# ---------- CLI ----------
def main() -> None:
    parser = argparse.ArgumentParser(description="Extract Arabic judgments from Word files to canonical JSON/JSONL")
    parser.add_argument("--input-dir", required=True, help="Directory containing Word documents")
    parser.add_argument("--output-dir", required=True, help="Directory for JSON outputs")
    parser.add_argument("--recursive", action="store_true", help="Scan input directory recursively")
    parser.add_argument("--export-labelstudio", action="store_true", help="Also export Label Studio JSONL tasks")
    args = parser.parse_args()

    input_dir = Path(args.input_dir)
    output_dir = Path(args.output_dir)

    if not input_dir.exists() or not input_dir.is_dir():
        raise SystemExit(f"Input directory does not exist or is not a directory: {input_dir}")

    output_dir.mkdir(parents=True, exist_ok=True)
    per_doc_dir = output_dir / "json"
    per_doc_dir.mkdir(parents=True, exist_ok=True)

    files = sorted(iter_input_files(input_dir, recursive=args.recursive), key=lambda p: str(p).lower())
    if not files:
        raise SystemExit(f"No supported Word files found in: {input_dir}")

    all_docs: List[dict] = []
    report_rows: List[dict] = []
    ls_tasks: List[dict] = []
    seen_doc_ids: dict[str, int] = {}
    seen_json_names: dict[str, int] = {}

    print(f"Found {len(files)} file(s). Starting extraction...", file=sys.stderr)

    for idx, path in enumerate(files, start=1):
        print(f"[{idx}/{len(files)}] {path.name}", file=sys.stderr)
        result = extract_document(path)

        base_doc_id = result.data["doc_id"]
        doc_id_count = seen_doc_ids.get(base_doc_id, 0) + 1
        seen_doc_ids[base_doc_id] = doc_id_count
        if doc_id_count > 1:
            unique_doc_id = f"{base_doc_id}__{doc_id_count}"
            result.data["doc_id"] = unique_doc_id
            result.report_row["doc_id"] = unique_doc_id
        else:
            unique_doc_id = base_doc_id

        all_docs.append(result.data)
        report_rows.append(result.report_row)

        base_json_name = slugify_filename(path.stem)
        json_name_count = seen_json_names.get(base_json_name, 0) + 1
        seen_json_names[base_json_name] = json_name_count
        out_name = f"{base_json_name}.json" if json_name_count == 1 else f"{base_json_name}__{json_name_count}.json"
        write_json(per_doc_dir / out_name, result.data)

        if args.export_labelstudio:
            ls_tasks.append(build_labelstudio_task(result.data))

    write_jsonl(output_dir / "judgments_master.jsonl", all_docs)
    write_csv(output_dir / "extraction_report.csv", report_rows)

    if args.export_labelstudio:
        write_jsonl(output_dir / "labelstudio_tasks.jsonl", ls_tasks)

    ok_count = sum(1 for r in report_rows if r["status"] == "ok")
    err_count = len(report_rows) - ok_count
    print(f"Done. OK={ok_count} ERROR={err_count}", file=sys.stderr)
    print(f"JSONL: {output_dir / 'judgments_master.jsonl'}", file=sys.stderr)
    print(f"Report: {output_dir / 'extraction_report.csv'}", file=sys.stderr)
    print(f"Per-doc JSON: {per_doc_dir}", file=sys.stderr)
    if args.export_labelstudio:
        print(f"Label Studio JSONL: {output_dir / 'labelstudio_tasks.jsonl'}", file=sys.stderr)


if __name__ == "__main__":
    main()
